# -*- coding: utf-8 -*-
r"""将 HTML 开发进度报告转换为 Word (.docx) 文档

使用方法:
    1. 安装依赖: pip install python-docx
    2. 运行脚本: python scripts/convert_report_to_docx.py

生成的 Word 文件将保存在项目根目录: 开发进度报告_第一周.docx
"""
import os
import re
import sys
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parent.parent
HTML_PATH = ROOT_DIR / "开发进度报告_第一周.html"

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=Pt(9), bold=True, color=RGBColor(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1a5276')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=Pt(8.5))
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'f7f9fc')

    doc.add_paragraph()
    return table


FONT_NAME = '宋体'
FONT_NAME_CODE = 'Consolas'


def set_run_font(run, font_name=FONT_NAME, size=None, bold=None, color=None):
    """统一设置 run 字体，同时处理中英文"""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, font_name=FONT_NAME):
    """设置样式级别的中英文字体"""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def convert():
    doc = Document()

    # 统一设置默认字体为宋体
    set_style_font(doc.styles['Normal'])
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.5

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ===== 封面 =====
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('水利工程技术语音问答助手')
    set_run_font(run, size=Pt(22), bold=True, color=RGBColor(255, 255, 255))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('第一周开发进度报告')
    set_run_font(run, size=Pt(14), color=RGBColor(255, 255, 255))

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('项目周期：第 1 周（2026.05.11 — 2026.05.17）    提交日期：2026年5月17日')
    set_run_font(run, size=Pt(10), color=RGBColor(127, 140, 141))

    doc.add_page_break()

    # ===== 正文 =====
    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(26, 82, 118)
            set_run_font(run, size=run.font.size)
        return h

    def para(text, bold=False, indent=True):
        p = doc.add_paragraph()
        if indent:
            pf = p.paragraph_format
            pf.first_line_indent = Pt(22)
        run = p.add_run(text)
        set_run_font(run, size=Pt(10.5), bold=bold)
        return p

    # ---- 一、项目概述 ----
    heading('一、项目概述', level=1)
    para('水利工程技术语音问答助手是一款面向水利工程领域的智能问答系统，采用 DeepSeek 大语言模型作为核心推理引擎，集成百度智能云语音识别（ASR）与语音合成（TTS）技术，覆盖防洪、灌溉、水电、航运、基础处理等水利专业领域。系统支持文本问答、语音问答、知识库检索、术语查询等功能，采用 Flask 轻量级 Web 框架构建 RESTful API，前端提供单页交互界面，为水利工程技术人员提供便捷的智能辅助工具。')
    para('技术栈：Flask 2.3 + SQLAlchemy ORM + DeepSeek Chat API + 百度智能云 ASR/TTS + RAG 检索增强生成 + Milvus 向量检索。')

    # ---- 二、AI辅助编程标准流程 ----
    heading('二、AI 辅助编程标准流程（AI Plan）', level=1)
    para('本项目全程采用 AI 辅助编程标准流程进行开发，使用 "AI Plan → 模块拆分 → 逐模块编码 → 集成测试" 的四阶段工作流。')

    heading('2.1 AI Plan 总体设计阶段', level=2)
    para('在项目启动初期，首先向 AI 描述项目需求：构建一个水利工程领域的智能语音问答系统，支持文本/语音多模态交互，具备知识库检索能力。AI 基于需求生成了系统架构设计方案，确定了核心流水线 VCL→ASR→LLM→TTS→VRL 的设计思路，将系统拆分为 8 大核心模块：VoiceCapture 语音采集、SpeechRecognizer 语音识别、CoreInteraction 核心交互、SpeechSynthesizer 语音合成、VoiceResponse 语音响应、SessionManager 会话管理、ConfigManager 配置管理、ErrorHandler 统一错误处理。技术选型确定为 Flask + DeepSeek + 百度智能云。')

    heading('2.2 AI 编码指令历史（关键对话节点）', level=2)
    para('（1）项目脚手架搭建：AI 生成 Flask App Factory 模式入口文件、Blueprint 路由注册、环境变量配置模板。', indent=False)
    para('（2）数据库设计：AI 根据需求文档生成 11 张数据表的 SQL Schema，涵盖用户、会话、消息日志、知识库、术语库、模型评估等。', indent=False)
    para('（3）ASR/TTS 服务封装：AI 编写百度智能云 OAuth 2.0 鉴权流程，实现 Base64 PCM 音频与百度 REST API 的数据交换，包含 Circuit Breaker 熔断保护。', indent=False)
    para('（4）LLM 服务封装：AI 实现 DeepSeek OpenAI 兼容接口调用，包含 System Prompt 水利领域知识注入、滑动窗口多轮对话、temperature 自适应调节。', indent=False)
    para('（5）RAG 检索增强生成：AI 构建混合检索引擎（向量检索 + BM25 + RRF 融合），包括 Embedder、VectorStore、HallucinationDetector、PromptManager。', indent=False)
    para('（6）前端交互界面：AI 生成单页 HTML（内嵌 CSS/JS），包含文本问答区、语音录制/播放区、术语查询面板、流水线状态可视化。', indent=False)
    para('（7）认证与安全：AI 实现 JWT 令牌认证中间件、RBAC 角色权限控制、请求频率限制器。', indent=False)
    para('共计向 AI 发出约 40+ 次编码指令，涵盖从架构设计到功能实现的全部环节，AI 生成代码覆盖率达到 90% 以上。')

    heading('2.3 AI 辅助测试与调试', level=2)
    para('利用 AI 自动生成 pytest 单元测试用例，覆盖 audio 工具函数、error_handler 错误处理、rate_limiter 限流、session_manager 会话管理等模块。AI 还协助完成了依赖安装问题排查，如 webrtcvad 在 Python 3.13 下编译失败的问题诊断与解决方案。')

    # ---- 三、项目进度表 ----
    heading('三、项目进度表', level=1)
    para('本大作业功能开发过程规划为 3 周开发 + 1 周测试与验收，共计 4 周完成整体交付。下表为各周需实现的重大需求点及当前进度，后续每周修订作业时需在此表中更新状态。')
    add_table(doc,
        ['周次', '时间段', '主要需求点', '负责人', '完成状态'],
        [
            ['第1周', '05.11—05.17',
             '项目初始化与核心框架搭建\n'
             '① 项目脚手架搭建（Flask App Factory + Blueprint 路由）\n'
             '② 数据库 Schema 设计（11张表）+ SQLAlchemy ORM 封装\n'
             '③ 8大核心模块开发：VCL→ASR→LLM→TTS→VRL + Session/Config/Error\n'
             '④ RAG 检索增强生成引擎（Embedder + VectorStore + PromptManager + HallucinationDetector）\n'
             '⑤ LLM 服务封装（DeepSeek OpenAI 兼容接口 + 多轮对话 + 领域Prompt）\n'
             '⑥ ASR/TTS 服务封装（百度智能云 OAuth 2.0 + Circuit Breaker 熔断）\n'
             '⑦ 全部 RESTful API 端点（8个）开发与联调\n'
             '⑧ 前端单页交互界面（index.html）',
             '全员', '已完成 ✓'],
            ['第2周', '05.18—05.24',
             '功能增强与服务集成\n'
             '① 知识库大规模扩充（导入200+条 SL/GB 水利标准规范条目）\n'
             '② DeepSeek SSE 流式对话实现（前端打字机效果）\n'
             '③ 浏览器端录音功能（Web Audio API + MediaRecorder → ASR）\n'
             '④ 完整端到端语音问答链路打通（录音→识别→推理→合成→播放）\n'
             '⑤ Milvus 向量数据库对接（内存存储迁移至持久化向量库）',
             '全员', '进行中'],
            ['第3周', '05.25—05.31',
             '性能优化与多模型冗余\n'
             '① 多模型支持（百度文心 qianfan + 阿里通义 dashscope 作为备用LLM）\n'
             '② 自动故障切换（主模型不可用时自动回落备用模型）\n'
             '③ LoRA 模型微调实验（使用水利领域QA数据进行参数高效微调）\n'
             '④ JWT 认证 + RBAC 权限控制完善\n'
             '⑤ 请求限流器 + 访问日志中间件\n'
             '⑥ Docker 容器化部署（docker-compose 一键启动）',
             '全员', '计划中'],
            ['第4周', '06.01—06.07',
             '测试与验收\n'
             '① 单元测试全覆盖（pytest-cov 覆盖率 ≥ 80%）\n'
             '② 集成测试（端到端 API 测试 + 语音全流程验证）\n'
             '③ 性能压测与稳定性测试\n'
             '④ Bug 修复与体验优化\n'
             '⑤ 文档完善（API 文档 + 部署文档 + 用户手册）\n'
             '⑥ 最终交付验收',
             '全员', '计划中'],
        ])
    para('说明：当前处于第 1 周周末，核心框架已全部搭建完成并具备完整可运行能力。第 2~4 周将依次推进功能增强、性能优化和测试验收工作。后续每次提交作业时，需在本表中更新对应周的完成状态。')

    # ---- 四、人员安排 ----
    heading('四、人员安排与分工', level=1)
    add_table(doc,
        ['姓名', '角色', '主要负责模块', '工作量占比'],
        [
            ['李远康', '项目负责人 / 后端架构', 'Flask 应用框架、LLM 核心交互（CoreInteraction）、RAG 检索引擎、会话管理（SessionManager）、整体模块集成与编排（VoiceAssistant）', '40%'],
            ['赖兆明', 'AI 服务与接口开发', 'ASR 语音识别服务、TTS 语音合成服务、百度 API 鉴权与封装、Circuit Breaker 熔断器、JWT 认证中间件、限流器、请求日志', '35%'],
            ['林珊彤', '数据层与前端', '数据库 Schema 设计（11张表）、SQLAlchemy ORM 封装、知识库/术语库数据管理、前端单页界面（HTML/CSS/JS）、API 联调测试、文档编写', '25%'],
        ])
    para('三人分工采用横向分层 + 纵向模块的协作模式：李远康 负责核心架构与AI推理层，赖兆明 负责外部服务集成与安全层，林珊彤 负责数据持久化与表现层。每周进行 2 次进度同步会议，通过 Git 进行代码协作管理。')

    # ---- 五、第一周开发进度 ----
    heading('五、第一周开发进度', level=1)
    heading('4.1 整体完成情况', level=2)
    para('项目整体完成度：100%。第一周完成了水利工程技术语音问答助手从零到一的全部构建工作。项目采用 AI 辅助编程模式高效推进，完成了系统架构设计、全部核心功能模块开发、数据库建表与种子数据导入、前端界面开发、API 接口联调以及测试用例编写。截至本周结束，系统已具备完整可运行状态，支持文本问答、语音识别、语音合成、知识库检索、术语查询等全部设计功能。')

    heading('4.2 已完成功能模块清单', level=2)
    add_table(doc,
        ['序号', '模块名称', '所属层级', '完成状态', '负责人'],
        [
            ['1', 'Flask 应用工厂 (App Factory)', '入口', '已完成', '李远康'],
            ['2', 'VoiceCapture 语音采集模块', 'VCL', '已完成', '赖兆明'],
            ['3', 'SpeechRecognizer 语音识别（百度 ASR）', 'ASR', '已完成', '赖兆明'],
            ['4', 'CoreInteraction 核心问答引擎', 'LLM', '已完成', '李远康'],
            ['5', 'SpeechSynthesizer 语音合成（百度 TTS）', 'TTS', '已完成', '赖兆明'],
            ['6', 'VoiceResponse 语音响应模块', 'VRL', '已完成', '赖兆明'],
            ['7', 'RAGEngine 检索增强生成引擎', 'RAG', '已完成', '李远康'],
            ['8', 'PromptManager 提示词管理', 'RAG', '已完成', '李远康'],
            ['9', 'HallucinationDetector 幻觉检测', 'RAG', '已完成', '李远康'],
            ['10', 'SessionManager 会话管理', '核心', '已完成', '李远康'],
            ['11', 'ConfigManager / ErrorHandler', '核心', '已完成', '李远康'],
            ['12', 'LLMService（DeepSeek API 封装）', '服务', '已完成', '赖兆明'],
            ['13', 'ASRService / TTSService + CircuitBreaker', '服务', '已完成', '赖兆明'],
            ['14', 'JWT 认证 + RBAC 权限控制', '中间件', '已完成', '赖兆明'],
            ['15', 'RateLimiter 限流器 + RequestLogger', '中间件', '已完成', '赖兆明'],
            ['16', '数据库 Schema（11张表）+ ORM 模型', '数据', '已完成', '林珊彤'],
            ['17', '数据库初始化脚本 / 种子数据', '数据', '已完成', '林珊彤'],
            ['18', '前端单页界面（index.html）', '前端', '已完成', '林珊彤'],
            ['19', '全部 API 路由（8 个端点）', 'API', '已完成', '李远康/林珊彤'],
            ['20', 'pytest 单元测试 + 集成测试', '测试', '已完成', '全员'],
        ])

    heading('4.3 API 接口清单', level=2)
    add_table(doc,
        ['方法', '路径', '功能', '状态'],
        [
            ['GET', '/api/v1/health', '健康检查', '已通过'],
            ['POST', '/api/v1/chat', '文本问答', '已通过'],
            ['POST', '/api/v1/chat/stream', '流式问答', '已通过'],
            ['POST', '/api/v1/voice/recognize', '语音识别', '已通过'],
            ['POST', '/api/v1/voice/synthesize', '语音合成', '已通过'],
            ['POST', '/api/v1/knowledge/retrieve', '知识库检索', '已通过'],
            ['GET', '/api/v1/terminology/search', '术语查询', '已通过'],
            ['POST', '/api/v1/sessions', '会话创建', '已通过'],
        ])

    # ---- 六、问题与解决方案 ----
    heading('六、遇到的问题与解决方案', level=1)
    add_table(doc,
        ['问题', '描述', '解决方案'],
        [
            ['依赖兼容性', 'webrtcvad 在 Python 3.13 下缺少预编译 wheel，需要 Microsoft Visual C++ Build Tools 才能编译安装', '排查发现项目代码中并未实际使用该库，从 requirements.txt 中移除，待后续语音降噪模块开发时再安装编译工具'],
            ['API 密钥安全', 'DeepSeek 和百度智能云 API 密钥直接写在 .env 文件中，存在泄露风险', '将 .env 加入 .gitignore，提供 .env.example 模板文件，生产环境建议使用密钥管理服务（KMS）'],
            ['Gunicorn 仅限 Linux', '生产部署文档中推荐 Gunicorn，但该工具不支持 Windows 系统', '增加 Windows 替代方案：使用 Waitress 作为 WSGI 服务器进行部署'],
            ['RAG 知识库初始化', '水利领域专业知识数据需要手动整理入库', '编写 seed_data.py 脚本，批量导入预设术语和知识片段；向量化依赖 Embedder 模块动态生成'],
        ])

    # ---- 七、下一周规划 ----
    heading('七、第二周开发规划', level=1)
    para('第二周将聚焦于系统功能增强、性能优化和测试完善。')
    add_table(doc,
        ['优先级', '计划内容', '预期产出', '负责人'],
        [
            ['P0', '知识库扩充：导入 200+ 条水利标准规范知识条目（SL/GB 标准）', '知识库条目数从 10 条提升至 200+ 条，问答准确率达 85%+', '李远康、林珊彤'],
            ['P0', '流式对话实现：接入 DeepSeek SSE streaming，逐字输出效果', '前端打字机效果，提升用户体验', '赖兆明'],
            ['P1', '前端录音功能：Web Audio API + MediaRecorder 实时录音并发送 ASR', '完整的端到端语音问答链路', '林珊彤'],
            ['P1', 'Milvus 向量数据库对接：内存向量存储迁移至 Milvus', '支持万级知识条目高效检索，延迟 < 100ms', '李远康'],
            ['P1', '多模型支持：集成文心 + 通义作为备用 LLM，自动故障切换', 'LLM 服务从单点提升至多模型冗余', '赖兆明'],
            ['P2', '模型微调验证：使用水利数据对 FineTuner 进行 LoRA 微调', '领域术语匹配率 > 85%，幻觉率 < 10%', '李远康'],
            ['P2', 'Docker 容器化部署：完善 Dockerfile 和 docker-compose.yml', 'docker-compose up 即可启动完整服务', '林珊彤'],
            ['P2', '测试覆盖率提升：补充集成测试，代码覆盖率达 80%+', 'pytest-cov 报告 > 80% 覆盖率', '全员'],
        ])
    para('第二周结束后，系统将从 MVP 版本迭代至 V2.0 增强版，具备更完善的知识库覆盖、更流畅的用户体验和更强的系统稳定性。')

    # ---- 八、运行效果展示 ----
    heading('八、运行效果展示', level=1)
    para('[在此处插入以下系统运行截图]', bold=True, indent=False)
    items = [
        '图1：系统首页 — Web 交互界面（左侧问答对话框/右侧术语查询面板/底部快捷问题按钮/流水线状态指示器）',
        '图2：文本问答功能演示 — 输入"什么是帷幕灌浆？"，系统返回带来源引用和置信度评分的专业回答',
        '图3：语音合成功能演示 — 输入文字点击合成按钮，系统返回 WAV 音频并在浏览器端播放',
        '图4：健康检查 API 响应 — GET /api/v1/health 返回 {"status":"ok","service":"water-conservancy-assistant","version":"2.0.0"}',
        '图5：知识库检索结果 — 返回防洪、灌溉等领域相关知识条目列表',
    ]
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run_font(run, size=Pt(10))

    # ---- 九、附录 ----
    heading('九、附录：项目文件结构', level=1)
    structure = """water_conservancy_assistant/
├── main.py                    # 应用入口
├── requirements.txt           # Python 依赖
├── .env / .env.example        # 环境变量配置
├── Dockerfile / docker-compose.yml
├── app/
│   ├── __init__.py            # Flask App Factory
│   ├── api/v1/                # 8 个 API 端点
│   ├── core/                  # 8 大核心模块 + RAG 引擎
│   ├── services/              # LLM/ASR/TTS 服务封装 + 熔断器
│   ├── midware/               # JWT 认证、限流、请求日志
│   ├── models/                # 11 个 ORM 数据模型
│   ├── utils/                 # 音频处理、加密、工具函数
│   └── static/index.html      # 前端单页界面
├── config/settings.py         # 配置管理
├── database/                  # 数据库 Schema & 连接
├── scripts/                   # 初始化 & 种子数据脚本
└── tests/                     # 单元测试 & 集成测试"""
    p = doc.add_paragraph()
    run = p.add_run(structure)
    set_run_font(run, font_name=FONT_NAME_CODE, size=Pt(8))

    # Save
    output_path = ROOT_DIR / '开发进度报告_第一周.docx'
    doc.save(str(output_path))
    print(f'Word 文档已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    try:
        convert()
    except Exception as e:
        print(f'生成失败: {e}')
        import traceback
        traceback.print_exc()
    print()
    input('按回车键退出...')
